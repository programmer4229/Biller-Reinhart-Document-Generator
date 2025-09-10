from flask import Flask, request, send_file, send_from_directory, jsonify
from flask_cors import CORS
from docx import Document
import os
import tempfile
import secrets
import hashlib
from datetime import datetime, timedelta

# Initialize Flask app
app = Flask(__name__)
CORS(app)

# Security Configuration
# CHANGE THIS PASSWORD - Make it something your employees will remember
COMPANY_PASSWORD = "BillerReinhart2025!"  # Change this to your desired password

# Store active sessions (in production, use Redis or database)
active_sessions = {}

def hash_password(password):
    """Simple password hashing"""
    return hashlib.sha256(password.encode()).hexdigest()

def generate_session_token():
    """Generate a secure session token"""
    return secrets.token_urlsafe(32)

def is_valid_session(token):
    """Check if session token is valid and not expired"""
    if not token or token not in active_sessions:
        return False
    
    session_data = active_sessions[token]
    # Session expires after 8 hours
    if datetime.now() > session_data['expires']:
        del active_sessions[token]
        return False
    
    return True

# Define paths
TEMPLATE_DIR = os.path.join(os.path.dirname(__file__), "templates")

# Look for React build folder
build_folder = None
current_dir = os.path.dirname(__file__)

possible_build_paths = [
    os.path.join(current_dir, '..', 'frontend', 'build'),
    os.path.join(current_dir, '..', 'frontend', 'dist'),
    os.path.join(current_dir, 'frontend', 'build'),
    os.path.join(current_dir, 'frontend', 'dist')
]

print(f"Looking for build folder...")
for path in possible_build_paths:
    abs_path = os.path.abspath(path)
    if os.path.exists(abs_path):
        build_folder = abs_path
        print(f"✅ Found React build folder at: {build_folder}")
        break

if build_folder and os.path.exists(os.path.join(build_folder, 'index.html')):
    app.static_folder = build_folder
    app.static_url_path = '/'
    print(f"✅ Flask configured to serve static files from: {build_folder}")
else:
    print("⚠️  Flask running in API-only mode")

# Authentication Routes
@app.route('/login', methods=['POST'])
def login():
    try:
        data = request.get_json()
        password = data.get('password', '')
        
        # Check password
        if password == COMPANY_PASSWORD:
            # Generate session token
            token = generate_session_token()
            
            # Store session (expires in 8 hours)
            active_sessions[token] = {
                'created': datetime.now(),
                'expires': datetime.now() + timedelta(hours=8),
                'ip': request.remote_addr
            }
            
            return jsonify({
                'success': True,
                'token': token,
                'message': 'Login successful'
            })
        else:
            return jsonify({
                'success': False,
                'message': 'Invalid password'
            }), 401
            
    except Exception as e:
        print(f"Login error: {e}")
        return jsonify({
            'success': False,
            'message': 'Login failed'
        }), 500

@app.route('/health')
def health():
    return {
        "status": "Backend is running",
        "build_folder": build_folder,
        "active_sessions": len(active_sessions),
        "templates_dir": TEMPLATE_DIR,
        "templates_exist": os.path.exists(TEMPLATE_DIR)
    }

# Serve React App
@app.route('/')
def serve_react():
    if build_folder and os.path.exists(os.path.join(build_folder, 'index.html')):
        return send_from_directory(build_folder, 'index.html')
    else:
        return """
        <h1>Engineering Doc Generator</h1>
        <p>React build not found. Check deployment logs.</p>
        <p><a href="/health">Health Check</a></p>
        """

@app.route('/<path:path>')
def serve_static_files(path):
    if build_folder and os.path.exists(build_folder):
        file_path = os.path.join(build_folder, path)
        if os.path.exists(file_path):
            return send_from_directory(build_folder, path)
        else:
            if os.path.exists(os.path.join(build_folder, 'index.html')):
                return send_from_directory(build_folder, 'index.html')
    return f"File not found: {path}", 404

# Document generation helper functions
def has_placeholders(text, replacements):
    """Check if text contains any placeholder tags that need replacement"""
    for key in replacements.keys():
        if f"{{{{{key}}}}}" in text:
            return True
    return False

def replace_text_in_paragraph(paragraph, replacements):
    """Only replace text if placeholders are found"""
    full_text = ''.join(run.text for run in paragraph.runs)
    
    if not has_placeholders(full_text, replacements):
        return
    
    for key, value in replacements.items():
        full_text = full_text.replace(f"{{{{{key}}}}}", value)
    
    for run in paragraph.runs:
        run.text = ''
    if paragraph.runs:
        paragraph.runs[0].text = full_text
    else:
        paragraph.add_run(full_text)

def replace_text_in_table(table, replacements):
    """Only replace text in table cells if placeholders are found"""
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                replace_text_in_paragraph(paragraph, replacements)

def replace_text_in_header_footer(section, replacements):
    """Preserve header/footer formatting by only replacing when necessary"""
    for paragraph in section.header.paragraphs:
        replace_text_in_paragraph(paragraph, replacements)
    
    for table in section.header.tables:
        replace_text_in_table(table, replacements)

    for paragraph in section.footer.paragraphs:
        replace_text_in_paragraph(paragraph, replacements)
    
    for table in section.footer.tables:
        replace_text_in_table(table, replacements)

# Protected Document Generation Route
@app.route("/generate", methods=["POST"])
def generate_doc():
    try:
        # Check authentication
        auth_token = request.form.get("auth_token")
        if not is_valid_session(auth_token):
            return jsonify({"error": "Unauthorized access"}), 401
        
        template_name = request.form.get("template_name")
        if not template_name:
            return {"error": "No template name provided"}, 400
            
        template_path = os.path.join(TEMPLATE_DIR, template_name)
        print(f"Authenticated user accessing template: {template_path}")

        if not os.path.exists(template_path):
            return {"error": f"Template not found: {template_name}"}, 404

        doc = Document(template_path)
        replacements = {key: value for key, value in request.form.items() 
                       if key not in ["template_name", "auth_token"]}

        # Replace text in document
        for paragraph in doc.paragraphs:
            replace_text_in_paragraph(paragraph, replacements)
        
        for table in doc.tables:
            replace_text_in_table(table, replacements)

        for section in doc.sections:
            replace_text_in_header_footer(section, replacements)

        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
            doc.save(tmp.name)
            tmp_path = tmp.name

        return send_file(tmp_path, as_attachment=True, download_name="customized.docx")
    
    except Exception as e:
        print(f"Error in generate_doc: {str(e)}")
        return {"error": "Document generation failed"}, 500

# Cleanup expired sessions periodically
@app.before_request
def cleanup_sessions():
    """Remove expired sessions"""
    current_time = datetime.now()
    expired_tokens = [token for token, data in active_sessions.items() 
                     if current_time > data['expires']]
    
    for token in expired_tokens:
        del active_sessions[token]

if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5050))
    print(f"🚀 Starting Flask app on port {port}")
    print(f"🔒 Password protection enabled")
    print(f"📁 Templates directory: {TEMPLATE_DIR}")
    print(f"🌐 Build folder: {build_folder or 'Not found'}")
    app.run(debug=False, host="0.0.0.0", port=port)